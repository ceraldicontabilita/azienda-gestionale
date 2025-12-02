"""
Contract Generator - Generazione contratti da template DOCX
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime
import os
from typing import Dict, Tuple


class ContractGenerator:
    """Generatore contratti di lavoro"""
    
    def __init__(self):
        self.templates_dir = '/home/claude/azienda-cloud/backend/templates/contracts'
        self.output_dir = '/home/claude/azienda-cloud/backend/uploads/contracts'
        
        os.makedirs(self.templates_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)
    
    async def generate_contract(
        self,
        template_name: str,
        employee_id: int,
        contract_data: dict,
        db
    ) -> Tuple[str, str]:
        """
        Genera contratto da template
        
        Returns:
            (docx_path, pdf_path)
        """
        
        # Recupera dati dipendente
        employee = await db.fetch_one("""
            SELECT * FROM employees WHERE id = :id
        """, {'id': employee_id})
        
        if not employee:
            raise ValueError("Dipendente non trovato")
        
        # Crea documento
        doc = self._create_base_contract(template_name, employee, contract_data)
        
        # Salva DOCX
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"contratto_{employee['cognome']}_{employee['nome']}_{timestamp}"
        
        docx_path = os.path.join(self.output_dir, f"{filename}.docx")
        doc.save(docx_path)
        
        # Converti in PDF
        pdf_path = self._convert_to_pdf(docx_path)
        
        return docx_path, pdf_path
    
    def _create_base_contract(
        self, 
        template_name: str, 
        employee: dict, 
        contract_data: dict
    ) -> Document:
        """Crea contratto base"""
        
        doc = Document()
        
        # Intestazione
        header = doc.add_paragraph()
        header.add_run('CERALDI GROUP S.R.L.\n').bold = True
        header.add_run('Piazza Nazionale, 46 - 80143 Napoli (NA)\n')
        header.add_run('P.IVA: 04523831214\n')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Titolo
        if template_name == 'contratto_determinato':
            title = doc.add_heading('CONTRATTO DI LAVORO A TEMPO DETERMINATO', 0)
        elif template_name == 'contratto_indeterminato':
            title = doc.add_heading('CONTRATTO DI LAVORO A TEMPO INDETERMINATO', 0)
        elif template_name == 'contratto_part_time':
            title = doc.add_heading('CONTRATTO DI LAVORO PART-TIME', 0)
        else:
            title = doc.add_heading('CONTRATTO DI LAVORO', 0)
        
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Parti
        doc.add_paragraph('Tra le seguenti parti:').bold = True
        doc.add_paragraph()
        
        # Datore di lavoro
        p = doc.add_paragraph()
        p.add_run('DATORE DI LAVORO\n').bold = True
        p.add_run('CERALDI GROUP S.R.L., con sede legale in Napoli, Piazza Nazionale 46, ')
        p.add_run('P.IVA 04523831214, in persona del legale rappresentante pro tempore, ')
        p.add_run('di seguito denominata "Azienda"\n\n')
        
        # E
        e_para = doc.add_paragraph()
        e_para.add_run('E\n\n').bold = True
        e_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lavoratore
        p = doc.add_paragraph()
        p.add_run('LAVORATORE\n').bold = True
        p.add_run(f"{employee['nome']} {employee['cognome']}, ")
        p.add_run(f"nato/a a {employee['luogo_nascita'] or '...'} ")
        p.add_run(f"il {employee['data_nascita'].strftime('%d/%m/%Y') if employee['data_nascita'] else '...'}, ")
        p.add_run(f"C.F. {employee['codice_fiscale']}, ")
        p.add_run(f"residente in {employee['indirizzo'] or '...'}, ")
        p.add_run(f"{employee['cap'] or ''} {employee['citta'] or ''} ({employee['provincia'] or ''}), ")
        p.add_run('di seguito denominato/a "Lavoratore"\n\n')
        
        # Premesso che
        doc.add_paragraph('PREMESSO CHE').bold = True
        doc.add_paragraph(
            'L\'Azienda opera nel settore della ristorazione e ha necessità di assumere '
            'personale qualificato per lo svolgimento delle proprie attività.'
        )
        doc.add_paragraph()
        
        # Si conviene e si stipula
        doc.add_paragraph('SI CONVIENE E SI STIPULA QUANTO SEGUE:').bold = True
        doc.add_paragraph()
        
        # Art. 1 - Assunzione
        self._add_article(doc, '1', 'ASSUNZIONE',
            f"L'Azienda assume il Lavoratore con la qualifica di {contract_data['mansione']} "
            f"- {contract_data['livello']}, "
            f"con decorrenza dal {contract_data['data_inizio'].strftime('%d/%m/%Y')}."
        )
        
        # Art. 2 - Durata
        if template_name == 'contratto_determinato':
            durata_text = (
                f"Il presente contratto ha durata determinata e scadrà il "
                f"{contract_data.get('data_fine', date.today()).strftime('%d/%m/%Y')}, "
                f"salvo proroga o trasformazione a tempo indeterminato."
            )
        else:
            durata_text = (
                "Il presente contratto ha durata indeterminata e potrà essere risolto "
                "da entrambe le parti con preavviso secondo quanto previsto dal CCNL applicabile."
            )
        
        self._add_article(doc, '2', 'DURATA', durata_text)
        
        # Art. 3 - Orario
        if contract_data['tipo_orario'] == 'part-time':
            orario_text = (
                f"Il Lavoratore è assunto con orario di lavoro PART-TIME pari al "
                f"{contract_data.get('percentuale_part_time', 50)}% dell'orario normale. "
                f"L'orario di lavoro settimanale sarà concordato tra le parti."
            )
        else:
            orario_text = (
                "Il Lavoratore è assunto con orario di lavoro FULL-TIME. "
                "L'orario di lavoro è fissato in 40 ore settimanali distribuite su 5/6 giorni, "
                "secondo le esigenze aziendali e nel rispetto del CCNL."
            )
        
        self._add_article(doc, '3', 'ORARIO DI LAVORO', orario_text)
        
        # Art. 4 - Retribuzione
        retr_text = (
            f"Il Lavoratore percepirà una retribuzione oraria di €{contract_data['retribuzione_oraria']:.2f}, "
            f"per un totale mensile lordo di circa €{float(contract_data['retribuzione_oraria']) * 160:.2f} "
            f"(per orario full-time). "
            f"La retribuzione sarà corrisposta entro il giorno 27 di ogni mese mediante bonifico bancario."
        )
        
        self._add_article(doc, '4', 'RETRIBUZIONE', retr_text)
        
        # Art. 5 - Ferie e Permessi
        self._add_article(doc, '5', 'FERIE E PERMESSI',
            "Il Lavoratore ha diritto a ferie retribuite nella misura prevista dal CCNL Turismo "
            "(26 giorni lavorativi annui) e a permessi retribuiti (ROL) secondo quanto stabilito "
            "dalla contrattazione collettiva."
        )
        
        # Art. 6 - Periodo di prova
        self._add_article(doc, '6', 'PERIODO DI PROVA',
            "È stabilito un periodo di prova di 30 giorni di effettivo lavoro, durante il quale "
            "ciascuna parte può recedere dal contratto senza preavviso."
        )
        
        # Art. 7 - CCNL applicabile
        self._add_article(doc, '7', 'CCNL APPLICABILE',
            "Al presente contratto si applica il CCNL del settore Turismo, pubblici esercizi, "
            "ristorazione collettiva e commerciale e servizi di catering."
        )
        
        # Art. 8 - Privacy
        self._add_article(doc, '8', 'TRATTAMENTO DATI PERSONALI',
            "Il Lavoratore prende atto dell'informativa sul trattamento dei dati personali "
            "ai sensi del Regolamento UE 2016/679 (GDPR) e autorizza l'Azienda al trattamento "
            "dei propri dati per le finalità connesse al rapporto di lavoro."
        )
        
        # Art. 9 - Sede di lavoro
        self._add_article(doc, '9', 'SEDE DI LAVORO',
            "La sede di lavoro è stabilita presso Piazza Nazionale 46, Napoli. "
            "L'Azienda si riserva la facoltà di modificare la sede di lavoro per "
            "comprovate esigenze organizzative, nel rispetto della normativa vigente."
        )
        
        # Art. 10 - Norme finali
        self._add_article(doc, '10', 'NORME FINALI',
            "Per tutto quanto non previsto dal presente contratto si fa espresso "
            "riferimento al CCNL applicabile e alle norme di legge vigenti."
        )
        
        doc.add_paragraph()
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
        
        # Firma
        p1 = doc.add_paragraph()
        p1.add_run(f"Napoli, {date.today().strftime('%d/%m/%Y')}")
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Tabella firme
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'IL DATORE DI LAVORO'
        table.rows[0].cells[1].text = 'IL LAVORATORE'
        
        table.rows[1].cells[0].text = '\nCERALDI GROUP S.R.L.\n\n'
        table.rows[1].cells[1].text = f'\n{employee["nome"]} {employee["cognome"]}\n\n'
        
        table.rows[2].cells[0].text = '\n____________________\n(Firma)'
        table.rows[2].cells[1].text = '\n____________________\n(Firma)'
        
        return doc
    
    def _add_article(self, doc: Document, num: str, title: str, text: str):
        """Aggiungi articolo al contratto"""
        
        p = doc.add_paragraph()
        p.add_run(f'Art. {num} - {title}\n').bold = True
        p.add_run(text)
        doc.add_paragraph()
    
    def _convert_to_pdf(self, docx_path: str) -> str:
        """Converti DOCX in PDF"""
        
        # Placeholder - richiede LibreOffice o altro converter
        # Per ora ritorna stesso path
        
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # TODO: Implementa conversione vera
        # subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_path])
        
        return pdf_path


# ============================================================================
# UTILITY - Crea template base
# ============================================================================

def create_default_templates():
    """Crea template contratti default"""
    
    templates_dir = '/home/claude/azienda-cloud/backend/templates/contracts'
    os.makedirs(templates_dir, exist_ok=True)
    
    # Template determinato
    # Template indeterminato  
    # Template part-time
    
    # I template vengono generati dinamicamente dal ContractGenerator
    
    print("✅ Template contratti configurati")


if __name__ == '__main__':
    create_default_templates()
